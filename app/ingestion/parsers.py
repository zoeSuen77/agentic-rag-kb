from __future__ import annotations

from html.parser import HTMLParser
from pathlib import Path

from app.indexing.schemas import RawDocument
from app.utils.hashing import stable_hash
from app.utils.text import normalize_whitespace


class _HTMLTextExtractor(HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self.parts: list[str] = []

    def handle_data(self, data: str) -> None:
        if data.strip():
            self.parts.append(data.strip())

    def text(self) -> str:
        return "\n".join(self.parts)


def parse_file(path: Path) -> RawDocument:
    suffix = path.suffix.lower()
    if suffix in {".txt", ".md", ".markdown"}:
        text = path.read_text(encoding="utf-8", errors="ignore")
    elif suffix in {".html", ".htm"}:
        parser = _HTMLTextExtractor()
        parser.feed(path.read_text(encoding="utf-8", errors="ignore"))
        text = parser.text()
    elif suffix == ".pdf":
        text = _parse_pdf(path)
    elif suffix == ".docx":
        text = _parse_docx(path)
    else:
        raise ValueError(f"Unsupported file extension: {suffix}")

    return RawDocument(
        doc_id=stable_hash({"path": str(path), "size": path.stat().st_size}, prefix="doc"),
        text=normalize_whitespace(text),
        source=str(path),
        metadata={"source": str(path), "filename": path.name, "extension": suffix},
    )


def _parse_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except Exception as exc:  # pragma: no cover
        raise RuntimeError("Install pypdf to parse PDF files") from exc
    reader = PdfReader(str(path))
    return "\n\n".join(page.extract_text() or "" for page in reader.pages)


def _parse_docx(path: Path) -> str:
    try:
        import docx
    except Exception as exc:  # pragma: no cover
        raise RuntimeError("Install python-docx to parse DOCX files") from exc
    document = docx.Document(str(path))
    return "\n".join(paragraph.text for paragraph in document.paragraphs)

