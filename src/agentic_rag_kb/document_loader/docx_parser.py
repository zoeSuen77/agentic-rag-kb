"""DOCX parser for enterprise Word technical documents.

This parser extracts headings, paragraphs, and tables as text while preserving a
simple representation that downstream chunking can use for section boundaries.
"""

from __future__ import annotations

from pathlib import Path

from agentic_rag_kb.document_loader.base import Document
from agentic_rag_kb.document_loader.cleaners import clean_document_text
from agentic_rag_kb.document_loader.ids import build_doc_id


class DocxDocumentParser:
    """Parse `.docx` files using python-docx."""

    supported_extensions = {".docx"}

    def load(self, path: Path) -> list[Document]:
        """Extract text from Word paragraphs and tables."""

        try:
            import docx
        except ImportError as exc:  # pragma: no cover
            raise RuntimeError("Install python-docx or `pip install -e '.[documents]'` to parse DOCX.") from exc

        document = docx.Document(str(path))
        title = _read_core_title(path, document)
        text = clean_document_text("\n".join([*_paragraph_lines(document), *_table_lines(document)]))
        return [
            Document(
                doc_id=build_doc_id(path, text),
                source_path=path,
                title=title,
                text=text,
                metadata={"file_type": "docx", "filename": path.name},
            )
        ]


def _read_core_title(path: Path, document) -> str:
    title = getattr(document.core_properties, "title", "") or ""
    return title.strip() or path.stem


def _paragraph_lines(document) -> list[str]:
    lines: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style_name = paragraph.style.name.lower() if paragraph.style else ""
        if style_name.startswith("heading"):
            level = _heading_level(style_name)
            lines.append(f"{'#' * level} {text}")
        else:
            lines.append(text)
    return lines


def _table_lines(document) -> list[str]:
    lines: list[str] = []
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if any(cells):
                lines.append(" | ".join(cells))
    return lines


def _heading_level(style_name: str) -> int:
    for token in style_name.split():
        if token.isdigit():
            return min(max(int(token), 1), 6)
    return 2

